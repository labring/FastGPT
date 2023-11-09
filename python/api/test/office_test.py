import requests
import pytest
from docx import Document
import os
from tempfile import NamedTemporaryFile
from docx.shared import Inches

image_path = os.path.join(os.path.dirname(__file__), "test.png")

# 定义一个函数来创建一个新的Word文档，并添加一个图片
def create_test_docx_with_image():
    # 使用临时文件来避免文件名冲突
    temp_file = NamedTemporaryFile(delete=False, suffix='.docx')
    # 创建一个文档对象
    doc = Document()
    # 添加一个段落
    doc.add_paragraph("This is a test document with an image.")
    # 添加一个图片，确保提供的图片路径是有效的
    doc.add_picture(image_path, width=Inches(1.25))  # 图片宽度设为1.25英寸
    # 保存文档到临时文件
    doc.save(temp_file.name)
    # 关闭临时文件
    temp_file.close()
    # 返回文件路径
    return temp_file.name

# 定义一个函数，它将创建并发送多个Word文档，并返回响应对象列表
def get_responses():
    responses = []
    # 创建并发送10个文档
    for _ in range(10):
        test_file_path = create_test_docx_with_image()
        with open(test_file_path, "rb") as f:
            files = {"file": (os.path.basename(test_file_path), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post("http://127.0.0.1:6010/extract_text/", files=files)
            responses.append(response)
        # 测试完成后删除文件
        os.unlink(test_file_path)
    return responses

# 使用pytest的parametrize装饰器测试所有响应
@pytest.mark.parametrize("response", get_responses())
def test_response(response):
    # 断言响应的状态码为200
    assert response.status_code == 200
    # 断言响应的内容类型是application/json
    assert "application/json" in response.headers["Content-Type"]
    # 断言响应的数据包含文本信息
    assert "text" in response.json()
